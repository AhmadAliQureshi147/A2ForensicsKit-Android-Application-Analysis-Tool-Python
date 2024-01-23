import sys
import os
import subprocess
from PyQt5.QtWidgets import QApplication, QMainWindow, QPushButton, QVBoxLayout, QWidget, QFileDialog, QTextEdit
from PyQt5.QtCore import QThread, pyqtSignal
from androguard.misc import AnalyzeAPK
from androguard.core.bytecodes.apk import APK
from html.parser import HTMLParser
import docx
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


# Utility class to strip HTML tags
class MLStripper(HTMLParser):
    def __init__(self):
        super().__init__()
        self.reset()
        self.strict = False
        self.convert_charrefs = True
        self.text = []

    def handle_data(self, d):
        self.text.append(d)

    def get_data(self):
        return ''.join(self.text)

def strip_tags(html):
    s = MLStripper()
    s.feed(html)
    return s.get_data()

class DecompilerThread(QThread):
    finished = pyqtSignal(str)

    def __init__(self, apkPath):
        super().__init__()
        self.apkPath = apkPath

    def run(self):
        try:
            apktool_path = 'D:\\Apktool\\apktool.bat'  # Update this path
            downloads_folder = os.path.join(os.path.expanduser('~'), 'Downloads')
            apk_name = os.path.basename(self.apkPath).split('.')[0]
            output_dir = os.path.join(downloads_folder, f"decompiled_{apk_name}")

            print(f"APK File: {self.apkPath}")
            print(f"Output Directory: {output_dir}")

            cmd = [apktool_path, 'd', self.apkPath, '-o', output_dir]
            result = subprocess.run(cmd, check=True, stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)

            output = "<b>Decompilation Output:</b>\n" + result.stdout
            if result.stderr:
                output += "\n<b>Decompilation Error:</b>\n" + result.stderr
            else:
                output += f"\n<b>Decompilation completed. Check the folder:</b> {output_dir}"

            print(output)  # Print the output for debugging

            self.finished.emit(output)
        except Exception as e:
            self.finished.emit(f"<b>Error in decompiling APK:</b> {str(e)}")


class AnalysisThread(QThread):
    finished = pyqtSignal(str)

    def __init__(self, apkPath):
        super().__init__()
        self.apkPath = apkPath

    def run(self):
        try:
            a, d, dx = AnalyzeAPK(self.apkPath)
            analysis_text = (
                f"<b>APK Analysis:</b>\n\n"
                f"<b>Package Name:</b> {a.get_package()}\n\n"
                f"<b>Main Activity:</b> {a.get_main_activity()}\n\n"
            )
            permissions = a.get_permissions()
            activities = a.get_activities()
            services = a.get_services()
            providers = a.get_providers()
            receivers = a.get_receivers()
            sdk_version = a.get_androidversion_code()

            analysis_text += (
                f"<b>Permissions:</b> {permissions}\n\n"
                f"<b>Activities:</b> {activities}\n\n"
                f"<b>Services:</b> {services}\n\n"
                f"<b>Providers:</b> {providers}\n\n"
                f"<b>Receivers:</b> {receivers}\n\n"
                f"<b>SDK Version:</b> {sdk_version}\n"
            )
            self.finished.emit(analysis_text)
        except Exception as e:
            self.finished.emit(f"<b>Error in analyzing APK:</b> {str(e)}")

class StaticVulnerabilityAnalysisThread(QThread):
    finished = pyqtSignal(str)

    def __init__(self, apkPath):
        super().__init__()
        self.apkPath = apkPath

    def run(self):
        try:
            apk = APK(self.apkPath)
            analysis_result = "<b>Static Vulnerability Analysis Results:</b>\n"
            permissions = apk.get_permissions()
            analysis_result += f"<b>Identified Permissions:</b> {', '.join(permissions)}\n"
            # Additional logic for vulnerability analysis goes here
            self.finished.emit(analysis_result)
        except Exception as e:
            self.finished.emit(f"<b>Error in static vulnerability analysis:</b> {str(e)}")

class DecompilerApp(QMainWindow):
    def __init__(self, frontPageUI=None):
        super().__init__()
        self.frontPageUI = frontPageUI
        self.filename = None
        self.analysisDone = False
        self.analysisResult = ""
        self.decompilationResult = ""
        self.staticVulnerabilityAnalysisResult = ""
        self.initUI()

    def initUI(self):
        self.setWindowTitle('A2 Forensics Kit')
        self.setGeometry(100, 100, 800, 600)
        self.setStyleSheet(
            "QMainWindow { background-color: #4F86E4; }"
            "QPushButton { background-color: #90A8D2; border-style: outset;"
            "border-width: 2px; border-radius: 10px; border-color: beige;"
            "font: bold 14px; min-width: 10em; padding: 6px; }"
            "QPushButton:hover { background-color: #384E90; }"
            "QTextEdit { background-color : #90A8D2; font: 14px; }"
        )

        layout = QVBoxLayout()

        self.btnOpen = QPushButton('Open APK', self)
        self.btnOpen.clicked.connect(self.openFileDialog)
        layout.addWidget(self.btnOpen)

        self.btnAnalyze = QPushButton('Analyze APK', self)
        self.btnAnalyze.clicked.connect(self.analyzeAPK)
        layout.addWidget(self.btnAnalyze)

        self.btnDecompile = QPushButton('Decompile APK', self)
        self.btnDecompile.clicked.connect(self.decompileAPK)
        layout.addWidget(self.btnDecompile)

        self.btnStaticVulnerabilityAnalysis = QPushButton('Static Vulnerability Analysis', self)
        self.btnStaticVulnerabilityAnalysis.clicked.connect(self.performStaticVulnerabilityAnalysis)
        layout.addWidget(self.btnStaticVulnerabilityAnalysis)

        self.btnGenerateReport = QPushButton('Generate Word Report', self)
        self.btnGenerateReport.clicked.connect(self.generateWordReport)
        layout.addWidget(self.btnGenerateReport)

        self.btnBack = QPushButton('Back', self)
        self.btnBack.clicked.connect(self.goBack)
        layout.addWidget(self.btnBack)

        self.textArea = QTextEdit(self)
        self.textArea.setReadOnly(True)
        layout.addWidget(self.textArea)

        centralWidget = QWidget()
        centralWidget.setLayout(layout)
        self.setCentralWidget(centralWidget)
        self.show()

    def goBack(self):
        self.close()
        if self.frontPageUI:
            self.frontPageUI.show()

    def openFileDialog(self):
        self.filename, _ = QFileDialog.getOpenFileName(self, "Open APK", "", "APK Files (*.apk)")
        if self.filename:
            self.textArea.append(f"<b>File selected:</b> {self.filename}")

    def decompileAPK(self):
        if self.filename:
            self.textArea.append("<b>Decompiling APK, please wait...</b>")
            self.thread = DecompilerThread(self.filename)
            self.thread.finished.connect(self.onDecompilationFinished)
            self.thread.start()
        else:
            self.textArea.append("<b>No APK file selected</b>")

    def analyzeAPK(self):
        if self.filename:
            self.textArea.append("<b>Analyzing APK, please wait...</b>")
            self.analysisThread = AnalysisThread(self.filename)
            self.analysisThread.finished.connect(self.onAnalysisFinished)
            self.analysisThread.start()
        else:
            self.textArea.append("<b>No APK file selected</b>")

    def performStaticVulnerabilityAnalysis(self):
        if self.filename:
            self.textArea.append("<b>Performing Static Vulnerability Analysis, please wait...</b>")
            self.staticVulnerabilityAnalysisThread = StaticVulnerabilityAnalysisThread(self.filename)
            self.staticVulnerabilityAnalysisThread.finished.connect(self.onStaticVulnerabilityAnalysisFinished)
            self.staticVulnerabilityAnalysisThread.start()
        else:
            self.textArea.append("<b>No APK file selected</b>")

    def onDecompilationFinished(self, result):
        self.decompilationResult = result
        self.textArea.append(result)

    def onAnalysisFinished(self, result):
        self.analysisResult = result
        self.textArea.append(result)
        self.analysisDone = True

    def onStaticVulnerabilityAnalysisFinished(self, result):
        self.staticVulnerabilityAnalysisResult = result
        self.textArea.append(result)

    def generateWordReport(self):
        if not self.filename:
            self.textArea.append("<b>No APK file selected</b>")
            return

        doc = docx.Document()
        doc.add_heading('APK Analysis Report', 0)

        # Decompilation Results
       # doc.add_heading('Decompilation Results:', level=1)
        p = doc.add_paragraph(strip_tags(self.decompilationResult))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Analysis Results
        doc.add_heading('Analysis Results:', level=1)
        p = doc.add_paragraph(strip_tags(self.analysisResult))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Static Vulnerability Analysis Results
        doc.add_heading('Static Vulnerability Analysis Results:', level=1)
        p = doc.add_paragraph(strip_tags(self.staticVulnerabilityAnalysisResult))
        p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

        # Save the document
        filename = f'Decompiler_Report_{os.path.basename(self.filename).split(".")[0]}.docx'
        filepath = os.path.join(os.path.expanduser('~'), 'Documents', filename)
        doc.save(filepath)

        self.textArea.append(f"<b>Report generated successfully:</b> {filepath}")

def main():
    app = QApplication(sys.argv)
    ex = DecompilerApp()
    sys.exit(app.exec_())

if __name__ == '__main__':
    main()
